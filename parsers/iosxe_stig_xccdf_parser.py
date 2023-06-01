#!python3

"""
    Will take the XCCDF format from IOS XE STIG to generate
    a word doc, JSON, and configuration file to fix issues.

    Word doc gives a table of what is wrong along with a listing
    from the STIG of severity and fix actions.
"""


import json
from docx import Document
import time
import yaml
import os


current_time = time.strftime("%Y%m%d_%H%M%S")
# script_run_path = f"{ os.path.expanduser('~') }/docker_ansible_stig/"
script_run_path = f"{ os.getcwd() }/"

transfer_dir = f"{ script_run_path }/transfer/"

input_file = f"{ transfer_dir }xccdf-results.xml"
parsed_file = f"{ transfer_dir }/iosxe_xccdf-results_parsed_{ current_time }.xml"
dictionary_output_file = (
    f"{ transfer_dir }/iosxe_xccdf-results_dictionary_{ current_time }.txt"
)
json_output_file = (
    f"{ transfer_dir }/iosxe_xccdf-results_dictionary_{ current_time }.json"
)
stig_directory = f"{ script_run_path }/playbooks/roles/iosxeSTIG/files/"
stig_files = [
    f"{ stig_directory }U_Cisco_IOS-XE_Router_NDM_STIG_V2R1_Manual-xccdf.xml",
    f"{ stig_directory }U_Cisco_IOS-XE_Router_RTR_STIG_V2R1_Manual-xccdf.xml",
]


output_doc_filename = f"{ transfer_dir }/iosxe_STIG_output_{ current_time }.docx"
config_file_output = f"{ transfer_dir }/iosxe_STIG_config_{ current_time }.txt"
yaml_stig_fixes_file = f"{ script_run_path }/config_var_files/iosxe_stig_fixes.yaml"


# The XCCDF file from the Ansible Playbok comes in with \n and
# \t.  To make this more uniform, this removes those placeholders and puts in real
# tabs and newlines for easier parsing later.
def create_new_xccdf(file_contents):
    formatted_output = file_contents.replace("\\n", "\n").replace("\\t", "\t")
    with open(parsed_file, "w") as output_file:
        output_file.write(formatted_output)


# creates a dictionary output and JSON output file for parsing later on if required
# basically reads in a file, sees if a line starts with an ending XML delimiter or not
# and looks for specific cdf lines to parse on.
def create_new_output():
    with open(parsed_file) as file:
        input_lines = file.readlines()
    rule_list = []
    result_list = []
    output_dict = {}
    for line in input_lines:
        line = line.strip()
        # skip all lines that start with ending delimiter
        if line.startswith("</"):
            continue
        # if starts with rule-result, then parse out
        # the rule on the line from character 28 to the
        # ninth character from the end.
        if line.startswith("<cdf:rule-result"):
            line_parts = line.split()
            rule = line_parts[1].split("=")
            rule_list.append(rule[1][27:-9])
        # parse out the results (either pass or fail)
        # from the line starting with result.  Starts on
        # character 13 and gets all from the 13th from the end
        elif line.startswith("<cdf:result>"):
            result_list.append(line[12:-13])
    # if the rules match the number of results, then we are good
    # to create the dictionary for making the output files
    # dictionary created uses the rule ID as the key and the
    # result as the value
    if len(rule_list) == len(result_list):
        for i in range(len(rule_list)):
            output_dict[rule_list[i]] = result_list[i]
    else:
        print(
            f"Something went wrong.  Please consult the { input_file }"
            + "for the latest and greatest file"
        )
    # create the dictionary and JSON files.
    # returns the dictionary of rules as the key with the
    # result as the value
    with open(dictionary_output_file, "w") as output_file:
        for key in output_dict.keys():
            output_string = f"{ key } = { output_dict[key] }\n"
            output_file.write(output_string)
    with open(json_output_file, "w") as output_file:
        json.dump(output_dict, output_file, indent=4)
    return output_dict


# Some severities get messed up due to length not being the same
# this fixes that so that makes it consistent
def fix_severity_wording(severity):
    if "high" in severity.lower():
        return "high"
    if "low" in severity.lower():
        return "low"
    if "medium" in severity.lower():
        return "medium"


# parses stigs from the XCCDF file to get fixes and IDs
def parse_stigs():
    rules_dict = {}
    for file in stig_files:
        with open(file) as input_file:
            input_string = input_file.read()
        # split all lines in the list on the "><" because that means the end
        # of a section for the XML
        input_lines = input_string.split("><")
        # counter is used here because of the interations that have to happen
        # once a line is found to contain certain text, then the counter
        # increments so that the next line goes with the previous
        counter = 0
        while counter < len(input_lines):
            # this indicates the start of a new rule-id from character 11 to
            # one from the end.  Counter is incremented by 1 then the rule title
            # is found, incremented by 2 for the rule severity, and by 1 again
            # for the rule long title
            if input_lines[counter].startswith("G"):
                rule_id = input_lines[counter][10:-1]
                counter += 1
                rule_title = input_lines[counter][6:-7]
                counter += 2
                rule_severity = fix_severity_wording(input_lines[counter][-7:-1])
                counter += 2
                rule_long_title = input_lines[counter][6:-7]
                counter += 1
                # if the line does not start with fixtext then increment
                while not input_lines[counter].startswith("fixtext "):
                    counter += 1
                # second counter is to get character counts used for
                # getting the rule fix ref.  This resets to 0 during each iteration
                second_counter = 0
                while not input_lines[counter][second_counter] == ">":
                    second_counter += 1
                second_counter += 1
                rule_fix_ref = input_lines[counter][second_counter:-9]
                # once this runs, then we have the files so we now need to create the
                # dictionary that will hold all of this for use later.  Key is the
                # rule ID and each value is another dictionary of values
                rules_dict[rule_id] = {
                    "rule_title": rule_title,
                    "rule_severity": rule_severity,
                    "rule_long_title": rule_long_title,
                    "rule_fix_ref": rule_fix_ref,
                }
            counter += 1
    # returns the final result dictionary for creating files
    return rules_dict


# makes the output word doc from the findings dictionary and results dictionary
def print_output(rules_dict, findings_dict):
    output_doc = Document()
    # create the title of the document
    output_doc.add_heading("STIG Findings", 0)
    # creates a table of findings using rule ID and finding
    output_doc.add_heading("Finding Table", level=1)
    table = output_doc.add_table(rows=1, cols=2, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Rule ID"
    hdr_cells[1].text = "Finding"
    # iterates through the findings dictionary to create executive table
    # each line is just ID and finding.  Once table is complete, then
    # add a page break
    for key in findings_dict.keys():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = findings_dict[key].upper()
    output_doc.add_page_break()
    # for each rule in the finding dictionary, if the rule is in the rules dictionary
    # then add to the rule with values to the doc.  IF a pass, just put that and move
    # on.  Else put in fix
    for key in findings_dict.keys():
        if key in rules_dict.keys():
            output_doc.add_heading(key, level=1)
            output_doc.add_paragraph(rules_dict[key]["rule_long_title"]).bold = True
            p = output_doc.add_paragraph("Finding on device: ")
            p.add_run(findings_dict[key].upper()).bold = True
            output_doc.add_paragraph(f"Rule Title: { rules_dict[key]['rule_title'] }")
            output_doc.add_paragraph(
                f"Rule Severity: { rules_dict[key]['rule_severity'] }"
            )
            if findings_dict[key].upper() != "PASS":
                output_doc.add_paragraph("Suggested Fix: ")
                output_doc.add_paragraph(rules_dict[key]["rule_fix_ref"])
    # output doc saves the word doc
    output_doc.save(output_doc_filename)


# Creates the config file for all that failed using the stig fixes file.
def generate_config_fix(findings_dict):
    with open(yaml_stig_fixes_file) as yaml_file:
        fix_dict = yaml.safe_load(yaml_file)
    with open(config_file_output, "w") as config_file:
        for key in findings_dict.keys():
            if findings_dict[key].upper() != "PASS" and key in fix_dict.keys():
                config_file.write(f"!{ key }")
                config_file.write("\n")
                config_file.write(f"{ fix_dict[key] }")
                config_file.write("\n\n")


def main():
    # remove leading b from file
    with open(input_file) as file:
        xccdf_file_contents = file.read()
    if xccdf_file_contents[0] == "b":
        new_xccdf_file_contents = xccdf_file_contents[1:]
    create_new_xccdf(new_xccdf_file_contents)
    findings_dict = create_new_output()
    rules_dict = parse_stigs()
    print_output(rules_dict, findings_dict)
    generate_config_fix(findings_dict)


if __name__ == "__main__":
    main()
